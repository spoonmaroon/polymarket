from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path("/Users/goon/polymarket/homelab-plan")
DOCX_PATH = OUT_DIR / "Home_Lab_Rack_Build_Plan.docx"
RACK_PNG_PATH = OUT_DIR / "rack_front_layout.png"


COLORS = {
    "black": "111315",
    "ink": "1F2933",
    "muted": "64748B",
    "blue": "1D4ED8",
    "cyan": "0891B2",
    "gray_fill": "EEF2F7",
    "dark_fill": "17202A",
    "table_header": "E8EEF5",
    "ok": "166534",
    "warn": "92400E",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: Inches) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, *, bold=False, size=11, color="000000") -> None:
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["blue"])
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    if level == 1:
        style_run(run, bold=True, size=16, color=COLORS["blue"])
    elif level == 2:
        style_run(run, bold=True, size=13, color=COLORS["cyan"])
    else:
        style_run(run, bold=True, size=12, color=COLORS["ink"])


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        style_run(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix) :])
        style_run(run)
    else:
        run = paragraph.add_run(text)
        style_run(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, COLORS["table_header"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_width(cell, Inches(widths[idx]))
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        style_run(run, bold=True, size=10, color=COLORS["ink"])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, Inches(widths[idx]))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            style_run(run, size=9.5)
    doc.add_paragraph()


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, font, fill, max_width=None) -> None:
    x1, y1, x2, y2 = box
    lines = []
    words = text.split()
    line = ""
    if max_width is None:
        max_width = x2 - x1 - 20
    for word in words:
        candidate = f"{line} {word}".strip()
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= max_width or not line:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + (len(lines) - 1) * 4
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=fill)
        y += bbox[3] - bbox[1] + 4


def draw_wrapped_text(draw: ImageDraw.ImageDraw, xy, text: str, font, fill, max_width: int, line_gap: int = 6) -> int:
    x, y = xy
    words = text.split()
    lines = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= max_width or not line:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((0, 0), line, font=font)
        y += bbox[3] - bbox[1] + line_gap
    return y


def create_rack_image() -> None:
    width, height = 1500, 2200
    img = Image.new("RGB", (width, height), "#F6F8FB")
    draw = ImageDraw.Draw(img)
    title = load_font(48, True)
    subtitle = load_font(24)
    label = load_font(22, True)
    small = load_font(19)
    device = load_font(26, True)
    device_small = load_font(19)

    draw.text((70, 46), "15U Home Lab Rack Front Layout", fill="#101418", font=title)
    draw.text(
        (72, 108),
        "Accurate to rack U positions and device heights. CX4170a builds require a deeper rack than 18 in.",
        fill="#556171",
        font=subtitle,
    )

    rack_x, rack_y = 300, 180
    rack_w, rack_h = 760, 1875
    u_h = rack_h / 15
    draw.rounded_rectangle(
        (rack_x - 52, rack_y - 48, rack_x + rack_w + 52, rack_y + rack_h + 48),
        radius=26,
        fill="#0C0F12",
        outline="#303842",
        width=6,
    )
    draw.rounded_rectangle(
        (rack_x - 26, rack_y - 24, rack_x + rack_w + 26, rack_y + rack_h + 24),
        radius=18,
        fill="#151A20",
        outline="#4B5563",
        width=4,
    )
    draw.rectangle((rack_x, rack_y, rack_x + rack_w, rack_y + rack_h), fill="#0F141A")

    def y_for_u(top_u: int, bottom_u: int) -> tuple[int, int]:
        # U15 is top row, U1 is bottom row.
        top_index = 15 - top_u
        bottom_index = 16 - bottom_u
        return int(rack_y + top_index * u_h), int(rack_y + bottom_index * u_h)

    # subtle glass reflection
    for offset in range(0, 120, 16):
        x = rack_x + rack_w - 170 + offset
        draw.line((x, rack_y + 10, x - 420, rack_y + rack_h - 10), fill="#1F2937", width=3)

    components = [
        (15, 15, "Patch Panel / Cable Brush", "#E5E7EB", "#111827"),
        (14, 14, "MikroTik CRS310-8G+2S+IN\n8x 2.5GbE + 2x 10G SFP+", "#F8FAFC", "#111827"),
        (13, 13, "Shelf: OPNsense Firewall + HP ProDesk", "#DDE6F3", "#111827"),
        (12, 12, "QNAP TS-433eU-US\n4-bay NAS", "#111827", "#F8FAFC"),
        (11, 11, "Vented Blank / Airflow Gap", "#2A3441", "#CBD5E1"),
        (10, 7, "PC 1 - Sliger CX4170a\nRyzen 7 7700X / RTX 5060 Ti\nUbuntu + Windows VM + Gaming + ML", "#111827", "#FFFFFF"),
        (6, 6, "Vented Blank / Airflow Gap", "#2A3441", "#CBD5E1"),
        (5, 2, "PC 2 - Sliger CX4170a\ni7-12700F / RTX 3060 Ti (+ optional 3060 Ti)\nDocker + Jellyfin + Trading Workers", "#111827", "#FFFFFF"),
        (1, 1, "Power / Vent\nUPS outside rack", "#1F2937", "#E5E7EB"),
    ]

    for top_u, bottom_u, text, fill, text_fill in components:
        y1, y2 = y_for_u(top_u, bottom_u)
        margin = 8 if top_u != bottom_u else 6
        draw.rounded_rectangle(
            (rack_x + margin, y1 + margin, rack_x + rack_w - margin, y2 - margin),
            radius=10,
            fill=fill,
            outline="#64748B" if fill != "#111827" else "#93C5FD",
            width=2,
        )
        if "PC " in text:
            # RGB fan accents and white cable lines for the aesthetic direction.
            for fan_idx in range(3):
                cx = rack_x + 78 + fan_idx * 82
                cy = y1 + 72
                draw.ellipse((cx - 28, cy - 28, cx + 28, cy + 28), outline="#60A5FA", width=4)
                draw.ellipse((cx - 17, cy - 17, cx + 17, cy + 17), outline="#F472B6", width=3)
            draw.line((rack_x + rack_w - 92, y1 + 35, rack_x + rack_w - 35, y1 + 90), fill="#F8FAFC", width=5)
            draw.line((rack_x + rack_w - 105, y1 + 58, rack_x + rack_w - 32, y1 + 130), fill="#E5E7EB", width=4)
        if top_u == 12:
            # four NAS trays
            tray_w = 96
            for i in range(4):
                tx = rack_x + 80 + i * (tray_w + 20)
                draw.rounded_rectangle((tx, y1 + 25, tx + tray_w, y2 - 25), radius=6, fill="#F3F4F6", outline="#CBD5E1", width=2)
                draw.rectangle((tx + 12, y1 + 44, tx + tray_w - 12, y1 + 52), fill="#111827")
            draw_wrapped_text(
                draw,
                (rack_x + 555, y1 + 42),
                "QNAP TS-433eU-US 4-bay NAS",
                device_small,
                "#F8FAFC",
                180,
                4,
            )
        else:
            font = device if (top_u - bottom_u + 1) >= 2 else device_small
            draw_centered_text(draw, (rack_x + 210, y1 + 8, rack_x + rack_w - 30, y2 - 8), text, font, text_fill)

    # U labels
    for u in range(15, 0, -1):
        y1, y2 = y_for_u(u, u)
        draw.text((rack_x - 125, y1 + (y2 - y1) / 2 - 14), f"U{u:02d}", fill="#111827", font=label)
        draw.line((rack_x - 10, y1, rack_x + rack_w + 10, y1), fill="#374151", width=1)
    draw.line((rack_x - 10, rack_y + rack_h, rack_x + rack_w + 10, rack_y + rack_h), fill="#374151", width=1)

    # side notes
    notes_x = 1100
    draw.rounded_rectangle((notes_x, 270, 1435, 820), radius=18, fill="#FFFFFF", outline="#CBD5E1", width=2)
    draw.text((notes_x + 24, 300), "Fit Notes", fill="#111827", font=load_font(30, True))
    notes = [
        "15U total height",
        "Two 4U compute PCs",
        "QNAP is storage only",
        "TS-433eU is 2.5GbE, not 10GbE",
        "UPS should stay outside light racks",
        "Measure GPU, cooler, PSU",
    ]
    y = 360
    for note in notes:
        draw.ellipse((notes_x + 28, y + 6, notes_x + 42, y + 20), fill="#2563EB")
        y_next = draw_wrapped_text(draw, (notes_x + 56, y), note, small, "#334155", 250, 3)
        y = max(y + 58, y_next + 14)

    draw.rounded_rectangle((notes_x, 900, 1435, 1210), radius=18, fill="#111827", outline="#334155", width=2)
    draw.text((notes_x + 24, 930), "Look", fill="#F8FAFC", font=load_font(30, True))
    look = ["black rack", "white cables", "white/gray switch", "subtle RGB fans", "clean blank panels"]
    y = 990
    for item in look:
        draw.text((notes_x + 34, y), f"- {item}", fill="#E5E7EB", font=small)
        y += 42

    img.save(RACK_PNG_PATH, quality=95)


def create_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, COLORS["blue"]),
        ("Heading 2", 13, COLORS["cyan"]),
        ("Heading 3", 12, COLORS["ink"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("Home Lab Rack Build Plan")
    style_run(title_run, bold=True, size=24, color=COLORS["black"])
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        "15U black/white rack concept for media, firewall, Ubuntu/Windows compute, and ML trading workloads"
    )
    style_run(subtitle_run, size=11, color=COLORS["muted"])

    add_heading(doc, "Decision Summary", 1)
    add_body(
        doc,
        "Use the QNAP TS-433eU-US as a compact 1U storage appliance, not as the main compute server. "
        "Use the two reused PCs in Sliger CX4170a cases for GPU compute, gaming, Docker services, and trading workloads. "
        "Use a dedicated OPNsense firewall box rather than putting routing on the NAS or a gaming PC.",
    )
    add_bullets(
        doc,
        [
            "Keep the NAS, firewall, and compute roles separate so failures and updates do not take down everything.",
            "Upgrade RAM and NVMe before buying more GPUs; the current CPUs/GPUs are already useful.",
            "If using CX4170a cases, buy a rack with real rear cable clearance. Avoid 18 inch deep racks for this exact case.",
            "Unraid is not needed if the QNAP is the NAS. Use Ubuntu Server or Proxmox on PC 2 unless you replace the QNAP with a DIY NAS.",
        ],
    )

    add_heading(doc, "Rack Layout", 1)
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_run = picture_paragraph.add_run()
    picture_run.add_picture(str(RACK_PNG_PATH), width=Inches(5.65))
    caption = doc.add_paragraph()
    caption_run = caption.add_run("Figure 1. Front rack layout. U positions are accurate; product faceplates are schematic.")
    style_run(caption_run, size=9, color=COLORS["muted"])
    doc.add_page_break()

    add_table(
        doc,
        ["Rack U", "Device", "Role / Notes"],
        [
            ["U15", "Patch panel / cable brush", "White/black cable dress, short patch cables, labels."],
            ["U14", "MikroTik CRS310-8G+2S+IN", "8x 2.5GbE RJ45 and 2x 10G SFP+ for the local core."],
            ["U13", "Shelf: OPNsense firewall + HP ProDesk", "Firewall appliance plus the mini PC for always-on services."],
            ["U12", "QNAP TS-433eU-US", "1U 4-bay NAS for media, backups, shared datasets."],
            ["U11", "Vented blank", "Airflow gap above the hot compute box."],
            ["U10-U7", "PC 1 in Sliger CX4170a", "Ubuntu desktop, Windows VM, gaming, primary ML experiments."],
            ["U6", "Vented blank", "Airflow gap between the two compute PCs."],
            ["U5-U2", "PC 2 in Sliger CX4170a", "Docker/Jellyfin/trading services and second GPU compute node."],
            ["U1", "Power / vent", "Keep the UPS outside the rack unless the rack load rating is comfortably high."],
        ],
        [0.85, 2.15, 3.5],
    )

    add_heading(doc, "Hardware Inventory And Upgrades", 1)
    add_table(
        doc,
        ["Device", "Current Hardware", "Role", "Upgrade / Add"],
        [
            [
                "PC 1",
                "Ryzen 7 7700X, ASUS PRIME B650M-A AX6 II micro-ATX, 32GB DDR5, 1TB SSD, RTX 5060 Ti",
                "Ubuntu Desktop, Windows VM, gaming, ML/dev",
                "Board fits CX4170a. Upgrade to 64GB DDR5 minimum; 128GB ideal if supported/stable. Add 2-4TB NVMe. Verify CPU cooler under 153mm for CX4170a.",
            ],
            [
                "PC 2",
                "i7-12700F, MSI PRO B660M-A CEC WIFI DDR4 (MS-7D37) micro-ATX, 16GB DDR4, 1TB SSD, RTX 3060 Ti; optional second 3060 Ti",
                "Ubuntu Server or Proxmox, Docker, Jellyfin, trading workers, secondary GPU jobs",
                "Board fits CX4170a. Upgrade to 64GB DDR4 using matched DIMMs. Add 2-4TB NVMe. Dual GPUs are questionable on this micro-ATX board; check slot spacing, PSU, and airflow first.",
            ],
            [
                "HP ProDesk 400 G",
                "i5-7500T, 28GB RAM, 480GB SSD",
                "Always-on low-power services",
                "Use as-is for AdGuard/Pi-hole, Tailscale, Uptime Kuma, and possibly Home Assistant.",
            ],
            [
                "QNAP TS-433eU-US",
                "1U 4-bay NAS, ARM Cortex-A55, 4GB non-upgradeable RAM, dual 2.5GbE",
                "Storage/media appliance",
                "Use 2x16TB NAS drives to start if budget allows; 12TB drives are acceptable but fill a 4-bay NAS faster.",
            ],
            [
                "Firewall appliance",
                "TBD",
                "Router, firewall, VLANs, VPN",
                "Add OPNsense box with 4x 2.5GbE ports. Protectli or Qotom-style hardware is fine.",
            ],
        ],
        [1.15, 1.8, 1.55, 2.0],
    )

    add_heading(doc, "Storage Plan", 1)
    add_body(
        doc,
        "For a 4-bay NAS, 16TB drives are the cleaner long-term choice. 12TB IronWolf drives are fine if the deal is good, "
        "but four slots disappear quickly when media, backups, datasets, and model artifacts all live on the same NAS.",
    )
    add_table(
        doc,
        ["Stage", "Drive Layout", "Usable Capacity", "Notes"],
        [
            ["Start", "2 x 16TB RAID 1", "~16TB", "Safe starting point; one-drive failure tolerance."],
            ["Expand", "3 x 16TB RAID 5", "~32TB", "More capacity, still one-drive failure tolerance."],
            ["Full", "4 x 16TB RAID 5", "~48TB", "Best capacity use for this NAS; keep external/offsite backups."],
            ["Safer full", "4 x 16TB RAID 6", "~32TB", "Two-drive failure tolerance, less usable space."],
            ["Existing drive", "2TB Seagate", "Do not add to main pool", "Use as scratch, transfer, or offline backup. It would bottleneck a large RAID group."],
        ],
        [0.9, 1.65, 1.4, 2.55],
    )

    add_heading(doc, "Networking And Security", 1)
    add_table(
        doc,
        ["Layer", "Recommended Part", "Why"],
        [
            ["Firewall", "OPNsense on a 4-port 2.5GbE appliance", "Stable router/firewall, VLANs, WireGuard, DNS rules."],
            ["Switch", "MikroTik CRS310-8G+2S+IN", "Good compact core: 8x 2.5GbE and 2x 10G SFP+."],
            ["PC links", "10GbE NICs + SFP+ DAC cables where needed", "Use the 10G ports for PC 1 and PC 2. The QNAP remains 2.5GbE."],
            ["DNS/adblock", "AdGuard Home or Pi-hole on HP ProDesk", "Network-wide ad/tracker filtering and local DNS names."],
            ["Remote access", "Tailscale or WireGuard", "No open NAS/admin ports to the public internet."],
        ],
        [1.1, 2.05, 3.35],
    )
    add_body(
        doc,
        "Suggested VLANs: Main, Servers, Media, Trading/ML, IoT, Guest, and Management. Keep trading secrets and live trading services away from media/download containers.",
    )

    add_heading(doc, "Software Stack", 1)
    add_table(
        doc,
        ["Machine", "Software", "Notes"],
        [
            ["Firewall", "OPNsense", "WAN/LAN routing, VLANs, DHCP, firewall rules, WireGuard."],
            ["QNAP NAS", "QTS, SMB shares, snapshots, QNAP backups", "Storage only. Avoid exposing QNAP directly to the internet."],
            ["PC 1", "Ubuntu Desktop, NVIDIA drivers, Docker, KVM/virt-manager, Windows 11 VM, Steam, Ollama", "Primary workstation and gaming/ML box."],
            ["PC 2", "Ubuntu Server or Proxmox, Docker Compose/Dockge, Jellyfin/Plex, trading containers", "If Proxmox feels too much, start with Ubuntu Server."],
            ["HP ProDesk", "Ubuntu Server, AdGuard Home/Pi-hole, Tailscale, Uptime Kuma", "Low-power always-on utility node."],
        ],
        [1.1, 2.35, 3.05],
    )

    add_heading(doc, "Missing Equipment Checklist", 1)
    add_bullets(
        doc,
        [
            "15U rack with enough depth and load capacity. CX4170a needs rear cable clearance; choose a deeper rack rather than the 18 inch VEVOR cabinet.",
            "2x Sliger CX4170a cases, matching rails, and rack hardware for the compute PCs.",
            "RAM upgrades: PC 1 needs DDR5; PC 2 needs DDR4. Target 64GB each, with 128GB on PC 1 if the budget and board support make sense.",
            "2-4TB NVMe SSD for PC 1 and 2-4TB NVMe SSD for PC 2.",
            "2x16TB NAS HDDs to start; later expand to 4x16TB. If using 12TB drives, plan for lower ceiling.",
            "OPNsense firewall appliance with 4x 2.5GbE ports.",
            "MikroTik CRS310 switch, 10GbE NICs for the PCs if needed, SFP+ DAC cables, and white Cat6A patch cables.",
            "Vented blanks, brush panel, shelf, labels, Velcro ties, cable combs, and controlled RGB/white cable accents.",
            "UPS: CyberPower CP1500PFCRM2U or similar, preferably outside the rack if the cabinet load rating is low.",
        ],
    )

    add_heading(doc, "Fit Checks Before Buying", 1)
    add_bullets(
        doc,
        [
            "Both known motherboards are micro-ATX, so both fit the CX4170a motherboard tray.",
            "PC 1 board: ASUS PRIME B650M-A AX6 II. Use DDR5 RAM only.",
            "PC 2 board: MSI PRO B660M-A CEC WIFI DDR4 (MS-7D37). Use DDR4 RAM only.",
            "Front-panel wiring: confirm each current case uses standard power/reset/LED connectors before transplanting.",
            "GPU exact model and length: CX4170a supports GPUs up to 375mm with front fans.",
            "GPU height including power cable: CX4170a supports up to 158mm including connectors.",
            "CPU cooler height: CX4170a air cooler limit is 153mm; use a compatible cooler or 360mm AIO.",
            "PSU length: keep at or under 190mm including cables.",
            "Rack usable depth: do not use an 18 inch deep cabinet with CX4170a; use a deeper rack or switch to CX4150a.",
            "Rack load rating: two 4U PCs plus NAS and shelves can approach light cabinet limits. Keep UPS outside unless the rack is rated for it.",
        ],
    )

    add_heading(doc, "Product Links", 1)
    links = [
        ("Sliger CX4170a case", "https://www.sliger.com/products/cx4170a"),
        ("ASUS PRIME B650M-A AX6 II specs", "https://www.asus.com/motherboards-components/motherboards/prime/prime-b650m-a-ax6-ii/techspec/"),
        ("MSI PRO B660M-A CEC WIFI DDR4 specs", "https://www.msi.com/Motherboard/PRO-B660M-A-CEC-WIFI-DDR4/Specification"),
        ("QNAP TS-433eU-US", "https://store.qnap.com/ts-433eu-us.html"),
        ("QNAP TS-433eU hardware specs", "https://www.qnap.com/en-us/product/ts-433eu/specs/hardware"),
        ("MikroTik CRS310-8G+2S+IN", "https://mikrotik.com/product/crs310_8g_2s_in"),
        ("Seagate IronWolf NAS drives", "https://www.seagate.com/products/nas-drives/ironwolf-hard-drive/"),
        ("CyberPower CP1500PFCRM2U UPS", "https://www.cyberpowersystems.com/product/ups/pfc-sinewave/cp1500pfcrm2u/"),
        ("OPNsense install docs", "https://docs.opnsense.org/manual/install.html"),
        ("Jellyfin Docker install", "https://jellyfin.org/docs/general/installation/container/"),
    ]
    for label_text, url in links:
        paragraph = doc.add_paragraph()
        add_hyperlink(paragraph, label_text, url)

    add_heading(doc, "Bottom Line", 1)
    add_body(
        doc,
        "This build should be treated as a 15U aesthetic compute rack with a separate NAS appliance. "
        "The clean path is QNAP for storage, OPNsense for firewall, PC 1 for interactive Ubuntu/Windows/gaming/ML, "
        "PC 2 for Docker/media/trading workers, and HP ProDesk for low-power utility services.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_rack_image()
    create_docx()
    print(DOCX_PATH)
    print(RACK_PNG_PATH)
